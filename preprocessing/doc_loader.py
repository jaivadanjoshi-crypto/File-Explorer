from typing import List

import docx
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

import config


def load_and_split_docx(uploaded_file) -> List[Document]:
    """Read a Streamlit-uploaded .docx file and split it into chunks."""
    document = docx.Document(uploaded_file)
    text = "\n".join(p.text for p in document.paragraphs if p.text.strip())

    doc = Document(page_content=text, metadata={"source": uploaded_file.name})

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=config.CHUNK_SIZE,
        chunk_overlap=config.CHUNK_OVERLAP,
    )
    return splitter.split_documents([doc])